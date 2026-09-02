"""Render a document plan as a .docx that Confluence can import cleanly.

Every construct here is one Confluence's Word importer understands: built-in
heading styles, plain paragraphs, real numbering definitions, bordered data
tables, and inline images. There are deliberately no text boxes, no floating
frames, and no layout tables -- those are exactly what turns an import into
phantom empty tables with orphaned text.
"""

from __future__ import annotations

import os
import re
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Emu, Pt, RGBColor

from .extract import RawDoc

BODY_FONT = "Arial"
BODY_SIZE_PT = 11

# 6.5" of text between 1" margins on US Letter.
MAX_IMAGE_WIDTH_PT = 468.0
EMU_PER_POINT = 12700

INDENT_STEP_TWIPS = 360
LIST_INDENT_TWIPS = 720
LIST_HANGING_TWIPS = 360

ORDERED_FORMATS = ["decimal", "lowerLetter", "lowerRoman", "decimal", "lowerLetter"]
BULLET_CHARS = ["•", "◦", "▪", "•", "◦"]


# --------------------------------------------------------------------------
# document chrome
# --------------------------------------------------------------------------

def _apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Word's stock heading styles are blue Calibri Light; recolour them so the
    # imported page doesn't arrive looking like a template.
    for level, size in ((1, 18), (2, 14), (3, 12)):
        try:
            style = document.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(16 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    try:
        title_style = document.styles["Title"]
        title_style.font.name = BODY_FONT
        title_style.font.size = Pt(26)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
    except KeyError:
        pass


def _ensure_hyperlink_style(document: Document) -> bool:
    try:
        document.styles["Hyperlink"]
        return True
    except KeyError:
        pass
    try:
        style = document.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
        style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        style.font.underline = True
        return True
    except Exception:
        return False


def _set_letter_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Emu(7772400)   # 8.5"
        section.page_height = Emu(10058400)  # 11"
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, attr, Emu(914400))  # 1"


# --------------------------------------------------------------------------
# numbering
# --------------------------------------------------------------------------

class NumberingRegistry:
    """Creates one Word numbering definition per logical list.

    Every list_id gets its own `w:abstractNum` *and* its own `w:num`. Word and
    LibreOffice restart a list per `w:num` when it carries a startOverride, but
    macOS Quick Look (the Finder preview) ignores that override and treats every
    num that shares an abstractNum as one continuous list, and other importers
    may do the same. With a shared abstractNum, a five-entry index made the
    first body list start at 6. A private abstractNum per list leaves nothing
    for a viewer to merge.
    """

    def __init__(self, document: Document):
        self._document = document
        self._numbering = self._numbering_element()
        self._by_list_id: dict[tuple[str, bool], int] = {}

    def _numbering_element(self):
        try:
            return self._document.part.numbering_part.element
        except (KeyError, AttributeError, NotImplementedError):
            return None

    def _next_id(self, tag: str, attr: str) -> int:
        existing = [
            int(el.get(qn(attr)))
            for el in self._numbering.findall(qn(tag))
            if el.get(qn(attr)) and el.get(qn(attr)).isdigit()
        ]
        return (max(existing) + 1) if existing else 1

    def _add_abstract(self, ordered: bool) -> int:
        abstract_id = self._next_id("w:abstractNum", "w:abstractNumId")
        levels = []
        for depth in range(5):
            indent = LIST_INDENT_TWIPS * (depth + 1)
            if ordered:
                fmt = ORDERED_FORMATS[depth]
                lvl_text = f"%{depth + 1}."
                rpr = ""
            else:
                fmt = "bullet"
                lvl_text = BULLET_CHARS[depth]
                rpr = (
                    '<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" '
                    'w:hint="default"/></w:rPr>'
                )
            levels.append(
                f'<w:lvl w:ilvl="{depth}">'
                f'<w:start w:val="1"/>'
                f'<w:numFmt w:val="{fmt}"/>'
                f'<w:lvlText w:val="{lvl_text}"/>'
                f'<w:lvlJc w:val="left"/>'
                f'<w:pPr><w:ind w:left="{indent}" w:hanging="{LIST_HANGING_TWIPS}"/></w:pPr>'
                f"{rpr}"
                f"</w:lvl>"
            )
        element = parse_xml(
            f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_id}">'
            f'<w:multiLevelType w:val="multilevel"/>'
            f'{"".join(levels)}'
            f"</w:abstractNum>"
        )

        # Schema order: every abstractNum must precede every num.
        anchor = self._numbering.findall(qn("w:abstractNum"))
        if anchor:
            anchor[-1].addnext(element)
        else:
            self._numbering.insert(0, element)
        return abstract_id

    def num_id_for(self, list_id: str, ordered: bool) -> int | None:
        if self._numbering is None:
            return None
        key = (list_id or "default", ordered)
        if key in self._by_list_id:
            return self._by_list_id[key]

        abstract_id = self._add_abstract(ordered)
        num_id = self._next_id("w:num", "w:numId")
        element = parse_xml(
            f'<w:num {nsdecls("w")} w:numId="{num_id}">'
            f'<w:abstractNumId w:val="{abstract_id}"/>'
            # Belt and braces for viewers that key on the num rather than the
            # abstractNum: an explicit restart at 1.
            f'<w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>'
            f"</w:num>"
        )
        self._numbering.append(element)
        self._by_list_id[key] = num_id
        return num_id


def _nested_list_id(
    list_id: str,
    level: int,
    previous: tuple[str, int] | None,
    aliases: dict[str, str],
) -> str:
    """Keep a nested item in the list of the item directly above it.

    Models sometimes hand sub-steps their own list_id. Word renders that
    tolerably because each such list happens to start at "a.", but the
    sub-steps are still a separate list: they no longer belong to the parent,
    and viewers or importers that map each list to its own block show them as
    a fresh top-level "1." An item at level > 0 that follows an item from a
    different list is treated as nested in that list, and every later item
    with the same id follows it there.
    """
    original = list_id
    list_id = aliases.get(original, original)
    if level > 0 and previous is not None and previous[0] != list_id:
        aliases[original] = previous[0]
        list_id = previous[0]
    return list_id


def _list_resumes(
    blocks: list[dict[str, Any]],
    index: int,
    previous: tuple[str, int] | None,
    aliases: dict[str, str],
) -> bool:
    """True when the next list item after blocks[index] continues the open list.

    Only images and paragraphs may sit in between; a heading, table, or
    divider ends the list. A nested item counts as continuing, because
    _nested_list_id folds it into the open list.
    """
    if previous is None:
        return False
    for later in blocks[index + 1:]:
        kind = later.get("type")
        if kind in ("image", "paragraph"):
            continue
        if kind != "list_item":
            return False
        level = max(0, min(4, int(later.get("level") or 0)))
        raw = str(later.get("list_id") or "default")
        return aliases.get(raw, raw) == previous[0] or level > 0
    return False


def _apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)

    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(LIST_INDENT_TWIPS * (level + 1)))
    indent.set(qn("w:hanging"), str(LIST_HANGING_TWIPS))
    p_pr.append(indent)


# --------------------------------------------------------------------------
# runs
# --------------------------------------------------------------------------

def _style_run(run, spec: dict[str, Any]) -> None:
    run.font.name = BODY_FONT
    if spec.get("bold"):
        run.bold = True
    if spec.get("italic"):
        run.italic = True
    if spec.get("underline"):
        run.underline = True
    color = (spec.get("color") or "").strip().lstrip("#")
    if len(color) == 6:
        try:
            run.font.color.rgb = RGBColor.from_string(color.upper())
        except ValueError:
            pass


def _add_hyperlink(paragraph, url: str, spec: dict[str, Any], has_style: bool) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    link_element = OxmlElement("w:hyperlink")
    link_element.set(qn("r:id"), r_id)

    run = paragraph.add_run(spec.get("text", ""))
    _style_run(run, spec)
    if has_style:
        try:
            run.style = paragraph.part.document.styles["Hyperlink"]
        except (KeyError, AttributeError):
            has_style = False
    if not has_style:
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.underline = True

    link_element.append(run._r)
    paragraph._p.append(link_element)


def _normalize_url(raw: str) -> str:
    """Return a linkable URL, or "" if this is not one.

    A model asked for the link behind `wiki.example.com` may hand back the
    bare hostname. Dropping it would silently lose a link the reader expects to
    click, so a plausible hostname gets a scheme instead.
    """
    url = (raw or "").strip()
    if not url:
        return ""
    if url.startswith(("http://", "https://", "mailto:")):
        return url
    if "@" in url and "/" not in url and " " not in url:
        return f"mailto:{url}"
    host = url.split("/", 1)[0]
    if " " not in url and "." in host and not host.endswith("."):
        return f"https://{url}"
    return ""


def _write_runs(paragraph, runs: list[dict[str, Any]], has_link_style: bool) -> None:
    for spec in runs or []:
        text = spec.get("text", "")
        if not text:
            continue
        url = _normalize_url(spec.get("link", ""))
        if url:
            _add_hyperlink(paragraph, url, spec, has_link_style)
        else:
            _style_run(paragraph.add_run(text), spec)


# --------------------------------------------------------------------------
# tables
# --------------------------------------------------------------------------

def _add_table(document: Document, header: list[str], rows: list[list[str]]) -> None:
    columns = max([len(header)] + [len(r) for r in rows]) if (header or rows) else 0
    if columns == 0:
        return

    body = [r for r in rows if any((c or "").strip() for c in r)]
    total_rows = (1 if header else 0) + len(body)
    if total_rows == 0:
        return

    table = document.add_table(rows=total_rows, cols=columns)
    try:
        table.style = document.styles["Table Grid"]
    except KeyError:
        _apply_manual_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    width_twips = int(MAX_IMAGE_WIDTH_PT * 20)
    column_width = Emu(int(width_twips / columns) * 635)

    def fill(cell, text: str, bold: bool) -> None:
        cell.width = column_width
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run((text or "").strip())
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE_PT - 1)
        run.bold = bold

    offset = 0
    if header:
        for index in range(columns):
            cell = table.rows[0].cells[index]
            fill(cell, header[index] if index < len(header) else "", True)
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="EFEFEF"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
        _repeat_header(table.rows[0])
        offset = 1

    for row_index, values in enumerate(body):
        for index in range(columns):
            fill(
                table.rows[row_index + offset].cells[index],
                values[index] if index < len(values) else "",
                False,
            )


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    tr_pr.append(header_flag)


def _apply_manual_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "999999")
        borders.append(element)
    table._tbl.tblPr.append(borders)


# --------------------------------------------------------------------------
# images
# --------------------------------------------------------------------------

def _add_image(document: Document, path: str, width_pt: float, alt: str, indent: int) -> bool:
    if not os.path.exists(path):
        return False

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    if indent:
        paragraph.paragraph_format.left_indent = Emu(int(indent * INDENT_STEP_TWIPS * 635))

    run = paragraph.add_run()
    try:
        shape = run.add_picture(path, width=Emu(int(min(width_pt, MAX_IMAGE_WIDTH_PT) * EMU_PER_POINT)))
    except Exception:
        paragraph._p.getparent().remove(paragraph._p)
        return False
    _set_alt(shape, alt)
    return True


def _append_image(paragraph, path: str, width_pt: float, alt: str, max_width_pt: float) -> bool:
    """Add a picture on its own line at the end of an existing paragraph.

    Used for a screenshot that belongs to a list item. In its own paragraph the
    picture would sit between two items of the list, and while Word keeps
    counting across that, Confluence's importer closes the list at any
    paragraph that is not a list item and starts the next step at 1 again.
    """
    if not os.path.exists(path):
        return False
    run = paragraph.add_run()
    run.add_break()
    try:
        shape = run.add_picture(path, width=Emu(int(min(width_pt, max_width_pt) * EMU_PER_POINT)))
    except Exception:
        run._r.getparent().remove(run._r)
        return False
    _set_alt(shape, alt)
    return True


def _set_alt(shape, alt: str) -> None:
    if not alt:
        return
    try:
        shape._inline.docPr.set("descr", alt)
    except Exception:
        pass


# --------------------------------------------------------------------------
# entry point
# --------------------------------------------------------------------------

def _drop_repeated_title(blocks: list[dict[str, Any]], title: str) -> list[dict[str, Any]]:
    """Remove a leading heading that just repeats the document title.

    The plan carries the title in its own field, and some models also emit it as
    the first heading, which renders the name of the document twice at the top.
    Enforcing it here means it cannot depend on which engine ran.
    """
    if not title or not blocks or blocks[0].get("type") != "heading":
        return blocks

    def squashed(text: str) -> str:
        return re.sub(r"\W+", "", text).casefold()

    first = "".join(run.get("text", "") for run in blocks[0].get("runs") or [])
    # Only the very first block can be the restated title; the same words
    # appearing later are a real heading about that subject.
    return blocks[1:] if squashed(first) == squashed(title) else blocks


def render(plan: dict[str, Any], doc: RawDoc, out_path: str) -> dict[str, Any]:
    document = Document()
    _set_letter_page(document)
    _apply_base_styles(document)
    has_link_style = _ensure_hyperlink_style(document)
    numbering = NumberingRegistry(document)

    stats = {"headings": 0, "paragraphs": 0, "list_items": 0, "tables": 0, "images": 0, "links": 0}

    title = (plan.get("title") or "").strip()
    if title:
        paragraph = document.add_paragraph(style="Title")
        run = paragraph.add_run(title)
        run.font.name = BODY_FONT

    blocks = _drop_repeated_title(plan.get("blocks") or [], title)

    # (list_id, level) of the last list item, so a nested item that arrived
    # with a stray list_id can be folded into the list above it. Headings and
    # dividers end a list; images and note paragraphs do not.
    previous_list: tuple[str, int] | None = None
    list_aliases: dict[str, str] = {}
    # The paragraph and level of the last list item, while nothing but its own
    # attachments have followed it. A screenshot after a step, or a note that
    # sits between two steps of one list, goes into that paragraph after a
    # line break instead of into a paragraph of its own. Word would keep
    # counting across a separate paragraph, but Confluence's importer closes
    # the list at any paragraph that is not a list item and restarts at 1.
    open_item: tuple[Any, int] | None = None

    for index, block in enumerate(blocks):
        kind = block.get("type")

        if kind == "heading":
            previous_list = None
            open_item = None
            level = max(1, min(3, int(block.get("level") or 1)))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _write_runs(paragraph, block.get("runs"), has_link_style)
            stats["headings"] += 1

        elif kind == "paragraph":
            runs = block.get("runs") or []
            if not any((r.get("text") or "").strip() for r in runs):
                continue
            if open_item is not None and _list_resumes(blocks, index, previous_list, list_aliases):
                paragraph = open_item[0]
                paragraph.add_run().add_break()
                _write_runs(paragraph, runs, has_link_style)
                stats["paragraphs"] += 1
                continue
            open_item = None
            paragraph = document.add_paragraph()
            indent = int(block.get("indent") or 0)
            if indent:
                paragraph.paragraph_format.left_indent = Emu(
                    int(min(indent, 4) * INDENT_STEP_TWIPS * 635)
                )
            if block.get("align") == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _write_runs(paragraph, runs, has_link_style)
            stats["paragraphs"] += 1

        elif kind == "list_item":
            runs = block.get("runs") or []
            if not any((r.get("text") or "").strip() for r in runs):
                continue
            paragraph = document.add_paragraph()
            ordered = bool(block.get("ordered", True))
            level = max(0, min(4, int(block.get("level") or 0)))
            list_id = _nested_list_id(
                str(block.get("list_id") or "default"), level, previous_list, list_aliases
            )
            previous_list = (list_id, level)
            num_id = numbering.num_id_for(list_id, ordered)
            if num_id is not None:
                _apply_numbering(paragraph, num_id, level)
            else:
                paragraph.style = document.styles["List Number" if ordered else "List Bullet"]
            _write_runs(paragraph, runs, has_link_style)
            stats["list_items"] += 1
            open_item = (paragraph, level)

        elif kind == "table":
            open_item = None
            _add_table(document, block.get("header") or [], block.get("rows") or [])
            stats["tables"] += 1

        elif kind == "image":
            candidate = doc.image_by_id(str(block.get("image_id") or ""))
            if not candidate:
                continue
            width = float(block.get("width_pt") or candidate.width_pt or 300)
            alt = str(block.get("alt") or "")
            if open_item is not None:
                paragraph, level = open_item
                # Keep the picture inside the step's text column.
                column = MAX_IMAGE_WIDTH_PT - LIST_INDENT_TWIPS * (level + 1) / 20
                if _append_image(paragraph, candidate.path, width, alt, column):
                    stats["images"] += 1
                continue
            if _add_image(
                document,
                candidate.path,
                width,
                alt,
                int(block.get("indent") or 0),
            ):
                stats["images"] += 1

        elif kind == "divider":
            previous_list = None
            open_item = None
            paragraph = document.add_paragraph()
            p_pr = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "CCCCCC")
            borders.append(bottom)
            p_pr.append(borders)

    stats["links"] = sum(
        1
        for block in plan.get("blocks") or []
        for run in (block.get("runs") or [])
        if _normalize_url(run.get("link", ""))
    )

    document.save(out_path)
    return stats
