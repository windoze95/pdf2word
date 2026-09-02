"""Smoke tests for the parts that must not silently regress.

Run with:  .venv/bin/python tests/test_pipeline.py
No model calls, no network -- everything here is deterministic.
"""

from __future__ import annotations

import os
import re
import sys
import tempfile
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pdf2word import planner_ai, render, schema  # noqa: E402
from pdf2word.extract import ImageCandidate, RawDoc  # noqa: E402

FAILURES: list[str] = []


def check(label: str, condition: bool, detail: str = "") -> None:
    if condition:
        print(f"  ok   {label}")
    else:
        print(f"  FAIL {label} {detail}")
        FAILURES.append(label)


def docx_xml(path: str) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def docx_part(path: str, name: str) -> str:
    with zipfile.ZipFile(path) as archive:
        try:
            return archive.read(name).decode("utf-8")
        except KeyError:
            return ""


# --------------------------------------------------------------------------

def test_strict_schema() -> None:
    print("\nstrict schema (OpenAI structured outputs)")
    strict = schema.strict_schema(schema.PLAN_SCHEMA)

    def every_object_ok(node) -> bool:
        if isinstance(node, dict):
            if node.get("type") == "object" and "properties" in node:
                if node.get("additionalProperties") is not False:
                    return False
                if set(node.get("required", [])) != set(node["properties"]):
                    return False
            return all(every_object_ok(v) for v in node.values())
        if isinstance(node, list):
            return all(every_object_ok(v) for v in node)
        return True

    check("every object requires all properties", every_object_ok(strict))
    check("optional fields become nullable", "anyOf" in strict["properties"]["title"])
    check(
        "required fields stay unwrapped",
        "anyOf" not in strict["properties"]["blocks"]["items"]["properties"]["type"],
    )
    check(
        "strip_nulls removes nulls",
        schema.strip_nulls({"a": 1, "b": None, "c": [{"d": None}]}) == {"a": 1, "c": [{}]},
    )


def test_carry_list_join() -> None:
    print("\nlist continuity across chunk boundaries")
    earlier = [
        {"type": "list_item", "list_id": "steps", "runs": [{"text": "one"}]},
        {"type": "list_item", "list_id": "steps", "runs": [{"text": "two"}]},
    ]
    incoming = [
        {"type": "list_item", "list_id": schema.CARRY_LIST_ID, "runs": [{"text": "three"}]},
        {"type": "paragraph", "runs": [{"text": "aside"}]},
        {"type": "list_item", "list_id": "other", "runs": [{"text": "fresh"}]},
    ]
    planner_ai._join_carried_lists(earlier, incoming)
    check("carried item joins the open list", incoming[0]["list_id"] == "steps")
    check("a later separate list is untouched", incoming[2]["list_id"] == "other")

    orphan = [{"type": "list_item", "list_id": schema.CARRY_LIST_ID, "runs": [{"text": "x"}]}]
    planner_ai._join_carried_lists([], orphan)
    check("no crash when nothing precedes it", orphan[0]["list_id"] == schema.CARRY_LIST_ID)


def test_url_normalisation() -> None:
    print("\nlink targets")
    cases = {
        "https://a.com/x": "https://a.com/x",
        "http://a.com": "http://a.com",
        "wiki.example.com": "https://wiki.example.com",
        "portal.example.com/path": "https://portal.example.com/path",
        "someone@example.com": "mailto:someone@example.com",
        "": "",
        "not a url": "",
        "See the portal": "",
    }
    for raw, expected in cases.items():
        got = render._normalize_url(raw)
        check(f"{raw!r} -> {got!r}", got == expected, f"expected {expected!r}")


def test_paste_extraction() -> None:
    print("\npasted content (OneNote-shaped HTML)")
    import base64
    import io

    from PIL import Image

    from pdf2word import extract_html

    buf = io.BytesIO()
    Image.new("RGB", (300, 120), "steelblue").save(buf, "PNG")
    data_uri = "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()

    html = (
        '<html xmlns:o="urn:schemas-microsoft-com:office:office">'
        "<head><style>p{margin:0}</style></head><body>"
        '<div style="position:absolute;left:48px;top:90px;width:624px">'
        '<p style="font-family:Calibri;font-size:16.0pt;font-weight:bold">A Heading</p>'
        '<p><span style="color:#C00000;font-weight:bold">NOTE:</span>'
        '<span style="font-family:Calibri;mso-fareast-font-family:Calibri"> body text</span></p>'
        '<p><a href="https://wiki.example.com">wiki.example.com</a></p>'
        f'<p><img width=300 height=120 src="{data_uri}" alt="Shot"></p>'
        '<p><img src="file:///C:/Temp/clip_image002.png"></p>'
        "<table><tr><td>Title</td><td>Code</td></tr></table>"
        "</div></body></html>"
    )

    doc = extract_html.extract_paste(html, tempfile.mkdtemp(), title="A Heading")
    skeleton = doc.pages[0].text

    check("marked as a paste", doc.source_kind == "paste")
    check("inline image recovered", len(doc.pages[0].images) == 1)
    check("dead file:// image counted", doc.unresolved_images == 1)
    check("missing image is flagged in the text", "could not be read" in skeleton)
    check(
        "link captured with its anchor text",
        doc.pages[0].links
        and doc.pages[0].links[0]["url"] == "https://wiki.example.com"
        and "wiki" in doc.pages[0].links[0]["anchor_text"],
    )

    # The point of the skeleton: meaning survives, styling noise does not.
    for noise in ("mso-", "Calibri", "position:absolute", "xmlns", "<style>", "width:624px"):
        check(f"strips {noise!r}", noise not in skeleton)
    for signal in ("font-weight:bold", "#C00000", "<table>", 'id="paste_i1"'):
        check(f"keeps {signal!r}", signal in skeleton)

    check("image size converted px->pt", round(doc.pages[0].images[0].width_pt) == 225)


def test_title_not_repeated() -> None:
    print("\nduplicated title guard")
    drop = render._drop_repeated_title

    heading = lambda t: {"type": "heading", "level": 1, "runs": [{"text": t}]}
    para = {"type": "paragraph", "runs": [{"text": "body"}]}

    check(
        "leading heading matching the title is dropped",
        drop([heading("Field Operations Handbook"), para], "Field Operations Handbook")
        == [para],
    )
    check(
        "match ignores case and punctuation",
        drop([heading("FIELD  OPERATIONS-HANDBOOK"), para], "Field Operations Handbook")
        == [para],
    )
    check(
        "a different leading heading is kept",
        len(drop([heading("Document Index"), para], "Field Operations Handbook")) == 2,
    )
    check(
        "the same text later in the document is kept",
        len(drop([para, heading("Field Operations Handbook")], "Field Operations Handbook"))
        == 2,
    )
    check("no title means no change", len(drop([heading("Anything"), para], "")) == 2)


def test_render() -> None:
    print("\nrendering a plan to .docx")
    plan = {
        "title": "Test Document",
        "blocks": [
            {"type": "heading", "level": 1, "runs": [{"text": "Section One"}]},
            {
                "type": "paragraph",
                "runs": [
                    {"text": "Warning: ", "bold": True, "color": "FF0000"},
                    {"text": "read this"},
                    {"text": " and see ", "italic": True},
                    {"text": "the portal", "link": "https://example.com/portal"},
                ],
            },
            {"type": "list_item", "list_id": "a", "ordered": True, "level": 0,
             "runs": [{"text": "first"}]},
            {"type": "list_item", "list_id": "a", "ordered": True, "level": 1,
             "runs": [{"text": "nested"}]},
            {"type": "list_item", "list_id": "b", "ordered": True, "level": 0,
             "runs": [{"text": "restarts"}]},
            {"type": "list_item", "list_id": "c", "ordered": False, "level": 0,
             "runs": [{"text": "bullet"}]},
            {"type": "table", "header": ["Title", "Code"], "rows": [["Manager", "MGR1"]]},
            {"type": "paragraph", "runs": [{"text": ""}]},
            {"type": "image", "image_id": "missing_id", "alt": "gone"},
        ],
    }
    doc = RawDoc(path="x.pdf", page_count=1, pages=[], metadata={}, workdir="")
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    stats = render.render(plan, doc, out)

    xml = docx_xml(out)
    check("file was written", os.path.exists(out) and os.path.getsize(out) > 5000)
    check("uses built-in heading styles", 'w:val="Heading1"' in xml)
    check("no layout tables: exactly one tbl", xml.count("<w:tbl>") == 1)
    check("no text boxes", "w:txbxContent" not in xml and "<v:shape" not in xml)
    check("no floating frames", "w:framePr" not in xml)
    check("table header repeats", "w:tblHeader" in xml)
    check("red run survives", 'w:val="FF0000"' in xml)
    check("hyperlink emitted", "<w:hyperlink" in xml)
    check("empty paragraph dropped", stats["paragraphs"] == 1, f"got {stats['paragraphs']}")
    check("unknown image id skipped", stats["images"] == 0)
    check("counts add up", stats["list_items"] == 4 and stats["tables"] == 1)

    numbering = docx_part(out, "word/numbering.xml")
    check("numbering part written", "<w:abstractNum" in numbering)
    check(
        "three lists get three numbering instances",
        numbering.count("<w:num ") >= 3,
        f"got {numbering.count('<w:num ')}",
    )
    check("bullet format defined", 'w:val="bullet"' in numbering)
    # macOS Quick Look merges every num that shares an abstractNum into one
    # running count, so each list must own its abstractNum outright.
    abstract_per_num = re.findall(
        r'<w:num [^>]*>.*?<w:abstractNumId w:val="(\d+)"/>', numbering, re.S
    )
    check(
        "every list owns its own abstractNum",
        len(abstract_per_num) == len(set(abstract_per_num)),
        f"abstractNum ids per num: {abstract_per_num}",
    )

    rels = docx_part(out, "word/_rels/document.xml.rels")
    check("hyperlink target is external", "example.com/portal" in rels)


def test_nested_list_ids() -> None:
    print("\nnested items stay in the parent's list")
    from docx import Document

    def item(list_id: str, level: int, ordered: bool = True) -> dict:
        return {"type": "list_item", "list_id": list_id, "level": level,
                "ordered": ordered, "runs": [{"text": f"{list_id}{level}"}]}

    plan = {
        "title": "",
        "blocks": [
            item("a", 0), item("b", 1), item("b", 1), item("a", 0), item("c", 1, ordered=False),
            item("a", 0),
            {"type": "heading", "level": 2, "runs": [{"text": "Next"}]},
            item("d", 0), item("d", 1),
        ],
    }
    doc = RawDoc(path="x.pdf", page_count=1, pages=[], metadata={}, workdir="")
    out = os.path.join(tempfile.mkdtemp(), "nested.docx")
    render.render(plan, doc, out)

    nums = []
    for paragraph in Document(out).paragraphs:
        num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
        if num_pr is not None:
            nums.append((paragraph.text, int(num_pr.numId.val), int(num_pr.ilvl.val)))
    by_text = {}
    for text, num_id, level in nums:
        by_text.setdefault(text, set()).add(num_id)

    a_num = by_text["a0"]
    check("top-level items share one list", len(a_num) == 1)
    check("stray-id sub-steps fold into the parent list", by_text["b1"] == a_num)
    check("sub-steps keep their nesting level",
          all(level == 1 for text, _, level in nums if text == "b1"))
    check("nested bullets get their own bullet list", by_text["c1"] != a_num)
    check("a list after a heading is a new list", by_text["d0"] != a_num)
    check("its own sub-steps stay with it", by_text["d1"] == by_text["d0"])


def test_image_placement() -> None:
    print("\nimage placement")
    from PIL import Image

    workdir = tempfile.mkdtemp()
    png = os.path.join(workdir, "shot.png")
    Image.new("RGB", (800, 400), "white").save(png)

    candidate = ImageCandidate(
        id="p01_i1", path=png, page=1, y=0, x0=0, x1=200,
        width_pt=900, height_pt=450, is_background=False,
    )

    class FakeDoc(RawDoc):
        def image_by_id(self, image_id):
            return candidate if image_id == "p01_i1" else None

    doc = FakeDoc(path="x.pdf", page_count=1, pages=[], metadata={}, workdir=workdir)
    plan = {"title": "", "blocks": [
        {"type": "image", "image_id": "p01_i1", "alt": "A screenshot", "width_pt": 900},
    ]}
    out = os.path.join(workdir, "img.docx")
    stats = render.render(plan, doc, out)

    check("image embedded", stats["images"] == 1)
    xml = docx_xml(out)
    check("alt text recorded", "A screenshot" in xml)

    # 468pt is the text width between 1" margins on Letter; 900pt must be capped.
    import re
    widths = [int(v) for v in re.findall(r'<wp:extent cx="(\d+)"', xml)]
    check(
        "oversized image capped to the text column",
        widths and widths[0] <= 468 * 12700 + 5,
        f"got {widths}",
    )


def test_list_survives_interruptions() -> None:
    print("\nimages and notes between steps stay inside the step")
    import re
    from docx import Document
    from PIL import Image

    workdir = tempfile.mkdtemp()
    png = os.path.join(workdir, "shot.png")
    Image.new("RGB", (800, 400), "white").save(png)
    candidate = ImageCandidate(
        id="p01_i1", path=png, page=1, y=0, x0=0, x1=200,
        width_pt=900, height_pt=450, is_background=False,
    )

    class FakeDoc(RawDoc):
        def image_by_id(self, image_id):
            return candidate if image_id == "p01_i1" else None

    def step(text: str) -> dict:
        return {"type": "list_item", "list_id": "steps", "level": 0, "ordered": True,
                "runs": [{"text": text}]}

    doc = FakeDoc(path="x.pdf", page_count=1, pages=[], metadata={}, workdir=workdir)
    plan = {"title": "", "blocks": [
        step("Click the funnel icon"),
        {"type": "image", "image_id": "p01_i1", "alt": "funnel"},
        step("Edit the filters"),
        {"type": "paragraph", "runs": [{"text": "Note: a date is required."}]},
        step("Save"),
        {"type": "paragraph", "runs": [{"text": "Escalate if it still fails."}]},
        {"type": "heading", "level": 2, "runs": [{"text": "Next"}]},
    ]}
    out = os.path.join(workdir, "steps.docx")
    stats = render.render(plan, doc, out)

    paragraphs = Document(out).paragraphs
    positions = [i for i, p in enumerate(paragraphs)
                 if p._p.pPr is not None and p._p.pPr.numPr is not None]
    steps = [paragraphs[i] for i in positions]
    check("three steps in one list", len(steps) == 3
          and len({int(p._p.pPr.numPr.numId.val) for p in steps}) == 1)
    check("no paragraph sits between the steps",
          positions == list(range(positions[0], positions[0] + 3)), f"at {positions}")
    check("screenshot lives inside the step before it",
          stats["images"] == 1 and "<w:drawing" in steps[0]._p.xml)
    check("note between steps lives inside the step before it", "Note: a date" in steps[1].text)
    plain = [p for p in paragraphs if p.text.strip() and "numPr" not in p._p.xml
             and not p.style.name.startswith("Heading")]
    check("remark after the last step stays its own paragraph",
          len(plain) == 1 and plain[0].text.startswith("Escalate"))
    widths = [int(v) for v in re.findall(r'<wp:extent cx="(\d+)"', docx_xml(out))]
    check("picture fits the step's text column (6.5in minus the 0.5in indent)",
          widths and widths[0] <= (468 - 36) * 12700 + 5, f"got {widths}")


if __name__ == "__main__":
    test_strict_schema()
    test_carry_list_join()
    test_nested_list_ids()
    test_url_normalisation()
    test_paste_extraction()
    test_title_not_repeated()
    test_render()
    test_image_placement()
    test_list_survives_interruptions()

    print()
    if FAILURES:
        print(f"{len(FAILURES)} check(s) failed: {', '.join(FAILURES)}")
        raise SystemExit(1)
    print("All checks passed.")
